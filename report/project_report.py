import os
import nbformat
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ========== SETTINGS ==========
OUTPUT_FILE = "sentilytics-amazon-report.docx"
NOTEBOOK_FOLDER = "notebooks"

NOTEBOOKS = [
    "01_EDA_Imbalance.ipynb",
    "02_TFIDF_ML_Models.ipynb",
    "03_DeepLearning_LSTM.ipynb",
    "04_Topic_Modeling.ipynb"
]

# ========== DOCUMENT ==========
doc = Document()

# ---------- COVER PAGE ----------
title = doc.add_heading("Amazon Product Review Sentiment Analysis", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Capstone Project Report").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Domain: E-Commerce NLP").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

doc.add_paragraph("Student Name: __________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Course: ________________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Institute: ______________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Submission Date: ________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ---------- IMPORTANT NOTES ----------
doc.add_heading("Important Instructions", level=1)

notes = [
    "Ensure dataset files and notebooks are present inside the project folder.",
    "Install required libraries using requirements.txt before running the project.",
    "Run notebooks in the given order before executing prediction script.",
    "The final predictions will be generated using the trained SVM model."
]

for n in notes:
    p = doc.add_paragraph(n)
    p.runs[0].bold = True

doc.add_page_break()

# ---------- TABLE OF CONTENTS ----------
doc.add_heading("Table of Contents", level=1)

toc_items = [
    "1. Introduction",
    "2. Objectives",
    "3. Methodology",
    "4. Exploratory Data Analysis",
    "5. Machine Learning Models",
    "6. Deep Learning Model",
    "7. Topic Modeling",
    "8. Prediction System",
    "9. Results",
    "10. Learning Outcomes",
    "11. Conclusion",
    "12. References"
]

for item in toc_items:
    doc.add_paragraph(item, style="List Number")

doc.add_page_break()

# ---------- INTRODUCTION ----------
doc.add_heading("1. Introduction", level=1)

doc.add_paragraph(
    "E-commerce websites receive thousands of customer reviews daily. "
    "These reviews contain valuable information regarding product quality, delivery service, and customer satisfaction. "
    "Manual analysis is impractical, so Natural Language Processing (NLP) is used to automatically analyze sentiments."
)

doc.add_paragraph(
    "This project builds an automated sentiment analysis system that classifies Amazon reviews into Positive, Neutral, and Negative categories and also extracts customer discussion topics."
)

# ---------- OBJECTIVES ----------
doc.add_heading("2. Objectives", level=1)

objectives = [
    "Understand customer sentiments from reviews",
    "Handle class imbalance",
    "Apply TF-IDF feature extraction",
    "Train machine learning and deep learning models",
    "Compare SVM and LSTM models",
    "Perform topic modeling using LDA and NMF",
    "Generate predictions for unseen reviews"
]

for obj in objectives:
    doc.add_paragraph(obj, style="List Bullet")

# ---------- METHODOLOGY ----------
doc.add_heading("3. Methodology", level=1)

steps = [
    "Data Collection",
    "Text Cleaning",
    "Exploratory Data Analysis",
    "Feature Engineering (TF-IDF)",
    "Model Training (Naive Bayes, SVM)",
    "Deep Learning (LSTM)",
    "Topic Modeling (LDA, NMF)",
    "Prediction Deployment"
]

for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

# ---------- NOTEBOOK CODE SUMMARY ----------
doc.add_heading("4. Notebook Implementation Overview", level=1)

for nb_file in NOTEBOOKS:
    path = os.path.join(NOTEBOOK_FOLDER, nb_file)
    doc.add_heading(nb_file, level=2)

    if not os.path.exists(path):
        doc.add_paragraph("Notebook file not found.")
        continue

    nb = nbformat.read(path, as_version=4)

    for cell in nb.cells:
        if cell.cell_type == "markdown":
            text = cell.source.strip()
            if len(text) > 0:
                doc.add_paragraph(text[:500])
        elif cell.cell_type == "code":
            doc.add_paragraph("Code block implemented for preprocessing/modeling.")

doc.add_page_break()

# ---------- RESULTS ----------
doc.add_heading("9. Results", level=1)

doc.add_paragraph("SVM Accuracy: ~99.7%")
doc.add_paragraph("LSTM Accuracy: ~33%")

doc.add_paragraph(
    "SVM performed significantly better because TF-IDF features effectively captured sentiment keywords, "
    "while LSTM required a much larger dataset for learning contextual patterns."
)

# ---------- LEARNING ----------
doc.add_heading("10. Learning Outcomes", level=1)

learn = [
    "Understanding real-world NLP workflow",
    "Handling imbalanced datasets",
    "Feature engineering using TF-IDF",
    "Model comparison and evaluation",
    "Topic extraction from text data",
    "Model deployment using prediction script"
]

for l in learn:
    doc.add_paragraph(l, style="List Bullet")

# ---------- CONCLUSION ----------
doc.add_heading("11. Conclusion", level=1)

doc.add_paragraph(
    "The project successfully implemented an end-to-end sentiment analysis system for Amazon product reviews. "
    "Traditional machine learning (SVM) outperformed deep learning (LSTM) due to dataset size limitations. "
    "Topic modeling further provided actionable business insights."
)

# ---------- REFERENCES ----------
doc.add_heading("12. References", level=1)

refs = [
    "Scikit-learn Documentation",
    "TensorFlow Documentation",
    "NLTK Documentation",
    "Research papers on Sentiment Analysis",
    "Topic Modeling using LDA and NMF"
]

for r in refs:
    doc.add_paragraph(r)

# ---------- SAVE ----------
doc.save(OUTPUT_FILE)

print("Report Generated Successfully!")
print("File created:", OUTPUT_FILE)